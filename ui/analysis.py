from concurrent.futures import ThreadPoolExecutor, as_completed
import io
import json
import queue
import threading
import time
from pandas import json_normalize
import streamlit as st
from docx import Document
from io import BytesIO
import pdfplumber
import pdfplumber
import queue
import threading
from datetime import datetime
import uuid
import pandas as pd
import concurrent.futures
import requests
from threading import Thread
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
from agents.ingredients_agent import run as run_ingredients
from agents.technology_agent import run as run_technology
from agents.benefits_agent import run as run_benefits
from agents.situations_agent import run as run_situations
from agents.motivations_agent import run as run_motivations
from agents.outcomes_agent import run as run_outcomes
from agents.product_generation_agent import run as run_product_generation
from functions import upload_to_azure, get_scraper_data


class AnalysisUI:
    def __init__(self, auth_service):
        self.auth = auth_service
        self.agents = {
            "IngredientsAgent": run_ingredients,
            "TechnologyAgent": run_technology,
            "BenefitsAgent": run_benefits,
            "SituationsAgent": run_situations,
            "MotivationsAgent": run_motivations,
            "OutcomesAgent": run_outcomes
        }

    def wizard_navigation(self):
        print(st.session_state.current_project["description"])
        steps = [
            {"title": "What We Know", "icon": "🔍"},
            {"title": "Digester", "icon": "⚙️"},
            {"title": "Results", "icon": "📊"},
            {"title": "Product Ideas", "icon": "💡"}
        ]
        
        st.markdown("<div class='wizard-navigation'>", unsafe_allow_html=True)
        cols = st.columns(len(steps))
        
        for i, step in enumerate(steps):
            with cols[i]:
                step_num = i + 1
                is_completed = step_num in st.session_state.completed_steps
                is_current = step_num == st.session_state.wizard_step
                
                # Status indicator
                status_icon = "✓" if is_completed else "➤" if is_current else "•"
                status_color = "#4CAF50" if is_completed else "#2196F3" if is_current else "#9E9E9E"
                bg_color = "#e8f5e9" if is_completed else "#e3f2fd" if is_current else "#f5f5f5"
                border_color = "#4CAF50" if is_completed else "#2196F3" if is_current else "#e0e0e0"
                
                # Make clickable if completed
                
                
                st.markdown(f"""
                <div class="wizard-step" style="border-color: {border_color}; background: {bg_color};">
                    <div style="display: flex; align-items: center; margin-bottom: 0.5rem;">
                        <span style="color: {status_color}; font-size: 1.5rem; margin-right: 0.5rem;">{status_icon}</span>
                        <h4 style="margin-bottom: 0;">Step {step_num}</h4>
                    </div>
                    <div style="display: flex; align-items: center; gap: 0.5rem;">
                        <span style="font-size: 1.2rem;">{step['icon']}</span>
                        <p style="margin-bottom: 0;">{step['title']}</p>
                    </div>
                </div>
                """, unsafe_allow_html=True)

                if is_completed:
                    if st.button(
                        f"Go to Step :-{step_num}",
                        key=f"wizard_step_{step_num}",
                        help=f"Go to {step['title']}",
                        use_container_width=True
                    ):
                        st.session_state.wizard_step = step_num
                        st.rerun()
    
        

    def extract_text_from_page(self,page):
        try:
            text = page.get_text("text")
            if text and text.strip():
                return text
            # Fallback to OCR
            pix = page.get_pixmap(dpi=300)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            ocr_text = pytesseract.image_to_string(image)
            return f"[OCR Page]\n{ocr_text.strip()}"
        except Exception as e:
            return f"[ERROR] Failed to extract page: {e}"


    def process_files(self, files):
        text = ""
        for f in files:
            try:
                file_size_mb = len(f.getvalue()) / (1024 * 1024)
                processing_container = st.empty()

                with processing_container.container():
                    st.status(f"🔍 Processing {f.name}", expanded=True)
                    progress_bar = st.progress(0)
                    progress_text = st.empty()

                    if file_size_mb > 3:
                        st.warning("This is a large file (>3MB). This may take a moment...")

                    progress_text.markdown(f"🔹 Detected file: {f.name} ({file_size_mb:.2f} MB)")
                    time.sleep(0.5)

                if f.name.endswith(".pdf"):
                    progress_text.markdown("Reading PDF file...")

                    doc = fitz.open(stream=f.getvalue(), filetype="pdf")
                    total_pages = len(doc)
                    page_results = [None] * total_pages
                    errors = []

                    def process_page(i):
                        try:
                            return i, self.extract_text_from_page(doc[i])
                        except Exception as e:
                            return i, f"[ERROR] {str(e)}"

                    with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
                        futures = [executor.submit(process_page, i) for i in range(total_pages)]
                        for count, future in enumerate(concurrent.futures.as_completed(futures), start=1):
                            i, result = future.result()
                            page_results[i] = result
                            if "[ERROR]" in result:
                                errors.append(f"Page {i+1}: {result}")
                            if count % 5 == 0 or count == total_pages:
                                progress = count / total_pages
                                progress_bar.progress(progress)
                                progress_text.markdown(f"Processed {count}/{total_pages} pages")

                    text = "\n\n".join(page_results)

                    if errors:
                        st.warning(f"⚠️ {len(errors)} pages failed to extract text. Example: {errors[0]}")

                    progress_bar.progress(1.0)
                    progress_text.success(f"✅ Completed processing {total_pages} pages")
                    time.sleep(0.5)
                elif f.name.endswith(".docx"):
                    progress_text.markdown("Reading DOCX file...")
                    try:
                        doc = Document(f)
                        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                        text += "\n\n".join(paragraphs)
                        progress_bar.progress(1.0)
                        progress_text.success("✅ Completed reading DOCX file")
                    except Exception as e:
                        st.error(f"❌ Error reading DOCX: {str(e)}")

                elif f.name.endswith(".txt"):
                    progress_text.markdown("Reading TXT file...")
                    try:
                        text += f.getvalue().decode("utf-8", errors="ignore")
                        progress_bar.progress(1.0)
                        progress_text.success("✅ Completed reading TXT file")
                    except Exception as e:
                       st.error(f"❌ Error reading TXT: {str(e)}")
            except Exception as e:
                st.error(f"❌ Error processing {f.name}: {str(e)}")
            finally:
                processing_container.empty()
        return text





    def save_files_to_azure(self, files, file_type):
        """Save files to Azure and return metadata"""
        file_metadata = []
        for file in files:
            file_url = upload_to_azure(file)
            if file_url:
                file_metadata.append({
                    "name": file.name,
                    "type": file_type,
                    "url": file_url,
                    "uploaded_at": datetime.now().isoformat()
                })
        return file_metadata



    def show_step1_content(self):
        st.subheader("🔍 What We Know")
        
        # Initialize session state if not exists
        if 'file_metadata' not in st.session_state:
            project_data = self.auth.get_project(st.session_state.current_project["_id"])
            st.session_state.file_metadata = project_data.get("file_metadata", {
                "technical": [],
                "marketing": []
            })
            st.session_state.social_media_data = project_data.get("social_media_data", None)
            st.session_state.research_result = project_data.get("research_result", "")
            st.session_state.last_hashtags = project_data.get("last_hashtags", [])
            st.session_state.rnd_files = []
            st.session_state.mkt_files = []

        # Function to validate file sizes immediately
        def validate_files(uploaded_files, file_type=""):
            if not uploaded_files:
                return True
            
            MAX_SINGLE_FILE = 7 * 1024 * 1024  # 2MB
            MAX_TOTAL_SIZE = 17 * 1024 * 1024  # 10MB
            
            # Check individual file sizes
            for file in uploaded_files:
                if file.size > MAX_SINGLE_FILE:
                    st.error(f"❌ {file_type} file '{file.name}' exceeds 2MB limit ({file.size/1024/1024:.2f}MB)")
                    return False
            
            # Check total size
            total_size = sum(file.size for file in uploaded_files)
            if total_size > MAX_TOTAL_SIZE:
                st.error(f"❌ Total {file_type} size exceeds 10MB limit ({total_size/1024/1024:.2f}MB)")
                return False
                
            return True

        # Technical Documents Section
        with st.expander("📚 In my possession - Technical Documents", expanded=True):
            # Display existing files
            if st.session_state.file_metadata.get("technical"):
                st.markdown("### Uploaded Technical Documents")
                for file in st.session_state.file_metadata["technical"]:
                    col1, col2, col3 = st.columns([4, 1, 1])
                    with col1:
                        st.markdown(f"📄 **{file['name']}**")
                        st.caption(f"Uploaded: {datetime.fromisoformat(file['uploaded_at']).strftime('%Y-%m-%d %H:%M')}")
                    with col2:
                        st.markdown(
                            f"""
                            <a href="{file['url']}" download target="_blank">
                                <button style="background-color: #f0f0f0; border: none; padding: 6px 12px; border-radius: 4px; cursor: pointer;">
                                    ⬇️ {file['name']}
                                </button>
                            </a>
                            """,
                            unsafe_allow_html=True
                        )
            
            col1, col2 = st.columns(2)
            with col2:
                st.markdown("**Private Technical**")
                private_tech = st.file_uploader(
                    "Upload confidential R&D documents",
                    type=["pdf", "docx", "txt"],
                    accept_multiple_files=True,
                    key="private_tech",
                    label_visibility="collapsed",
                    on_change=lambda: validate_files(st.session_state.private_tech, "private technical")
                )
            with col1:
                st.markdown("**Public Technical**")
                public_tech = st.file_uploader(
                    "Upload published technical documents",
                    type=["pdf", "docx", "txt"],
                    accept_multiple_files=True,
                    key="public_tech",
                    label_visibility="collapsed",
                    on_change=lambda: validate_files(st.session_state.public_tech, "public technical")
                )
            
            # Validate and store files
            tech_files = []
            if private_tech and validate_files(private_tech, "private technical"):
                tech_files.extend(private_tech)
            if public_tech and validate_files(public_tech, "public technical"):
                tech_files.extend(public_tech)
            st.session_state.rnd_files = tech_files

        # Marketing Documents Section
        with st.expander("📊 In my possession - Marketing Documents", expanded=True):
            # Display existing files
            if st.session_state.file_metadata.get("marketing"):
                st.markdown("### Uploaded Marketing Documents")
                for file in st.session_state.file_metadata["marketing"]:
                    col1, col2, col3 = st.columns([4, 1, 1])
                    with col1:
                        st.markdown(f"📄 **{file['name']}**")
                        st.caption(f"Uploaded: {datetime.fromisoformat(file['uploaded_at']).strftime('%Y-%m-%d %H:%M')}")
                    with col2:
                        st.markdown(
                            f"""
                            <a href="{file['url']}" download target="_blank">
                                <button style="background-color: #f0f0f0; border: none; padding: 6px 12px; border-radius: 4px; cursor: pointer;">
                                    ⬇️ {file['name']}
                                </button>
                            </a>
                            """,
                            unsafe_allow_html=True
                        )
            
            col1, col2 = st.columns(2)
            with col2:
                st.markdown("**Private Marketing**")
                private_mkt = st.file_uploader(
                    "Upload internal marketing materials",
                    type=["pdf", "docx", "txt"],
                    accept_multiple_files=True,
                    key="private_mkt",
                    label_visibility="collapsed",
                    on_change=lambda: validate_files(st.session_state.private_mkt, "private marketing")
                )
            with col1:
                st.markdown("**Public Marketing**")
                public_mkt = st.file_uploader(
                    "Upload public marketing materials",
                    type=["pdf", "docx", "txt"],
                    accept_multiple_files=True,
                    key="public_mkt",
                    label_visibility="collapsed",
                    on_change=lambda: validate_files(st.session_state.public_mkt, "public marketing")
                )
            
            # Validate and store files
            mkt_files = []
            if private_mkt and validate_files(private_mkt, "private marketing"):
                mkt_files.extend(private_mkt)
            if public_mkt and validate_files(public_mkt, "public marketing"):
                mkt_files.extend(public_mkt)
            st.session_state.mkt_files = mkt_files

        # # Social Media Scraper Section

        # Initialize session state
        if 'hashtags_list' not in st.session_state:
            st.session_state.hashtags_list = []
        if 'hashtag_input' not in st.session_state:
            st.session_state.hashtag_input = ""
        if 'notification' not in st.session_state:
            st.session_state.notification = None
        if 'notification_time' not in st.session_state:
            st.session_state.notification_time = None

        def add_hashtag():
            cleaned_hashtag = st.session_state.hashtag_input.strip().replace("#", "").lower()
            
            
            if cleaned_hashtag in [h.lower() for h in st.session_state.hashtags_list]:
                show_notification("warning", f"#{cleaned_hashtag} already exists!")
                st.session_state.hashtag_input = ""
                return
            
            if len(st.session_state.hashtags_list) > 3:
                show_notification("error", "Maximum 3 hashtags allowed")
                return
            
            st.session_state.hashtags_list.append(cleaned_hashtag)
            st.session_state.hashtag_input = ""
            show_notification("success", f"Added #{cleaned_hashtag}")

        def remove_hashtag(index):
            removed = st.session_state.hashtags_list.pop(index)
            show_notification("info", f"Removed #{removed}")

        def show_notification(type, message):
            st.session_state.notification = {"type": type, "message": message}
            
        def clear_notification_if_rerun():
            # If it exists and wasn't just set this run, clear it
            if st.session_state.get("notification") and not st.session_state.get("notification_just_set"):
                st.session_state.notification = None
                st.session_state.notification_just_set = False  # reset the flag
            elif st.session_state.get("notification_just_set"):
                st.session_state.notification_just_set = False  # reset for next run

        # UI Layout
        with st.expander("📱 Social Media Scraper", expanded=True):
            # Display notifications
            if st.session_state.notification:
                if st.session_state.notification["type"] == "error":
                    st.error(st.session_state.notification["message"])
                elif st.session_state.notification["type"] == "warning":
                    st.warning(st.session_state.notification["message"])
                elif st.session_state.notification["type"] == "success":
                    st.success(st.session_state.notification["message"])
                else:
                    st.info(st.session_state.notification["message"])
            
            col1, col2 = st.columns([4, 1])
            with col1:
                st.text_input(
                    "Add a hashtag (max 3):",
                    placeholder="cleanbeauty",
                    key="hashtag_input"
                )
            
            with col2:
                st.write("")  # Spacer
                st.button("➕ Add", key="add_hashtag", on_click=add_hashtag)
            
            # Display current hashtags
            if st.session_state.hashtags_list:
                st.write("Current hashtags:")
                for idx, hashtag in enumerate(st.session_state.hashtags_list):
                    cols = st.columns([3, 1])
                    with cols[0]:
                        st.write(f"#{hashtag}")
                    with cols[1]:
                        st.button(
                            "Remove", 
                            key=f"remove_{idx}",
                            on_click=remove_hashtag,
                            args=(idx,)
                        )
    
            # Always run this early in app
            clear_notification_if_rerun()
            if st.button("🌐 Scrape 🎵 TikTok", key="scrape_button"):
                if not st.session_state.hashtags_list:
                    st.error("Please add at least one hashtag")
                else:
                    # Create progress elements
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    status_text.info(f"🎵 Starting TikTok scraping for {len(st.session_state.hashtags_list)} hashtags...")
                    
                    # Initialize progress tracking
                    update_queue = queue.Queue()
                    result_holder = {"data": {}}
                    error_holder = {"errors": []}
                    
                    # Make a local copy of hashtags
                    hashtags_to_process = st.session_state.hashtags_list.copy()
                    total_hashtags = len(hashtags_to_process)
                    
                    # Track progress per hashtag
                    progress_tracker = {hashtag: {"completed": 0, "total": 0} for hashtag in hashtags_to_process}
                    
                    def run_scraper(hashtag):
                        try:
                            def progress_callback(completed, total):
                                # Update progress for this hashtag
                                progress_tracker[hashtag]["completed"] = completed
                                progress_tracker[hashtag]["total"] = total
                                
                                # Calculate overall progress
                                total_completed = sum(p["completed"] for p in progress_tracker.values())
                                total_posts = sum(p["total"] for p in progress_tracker.values())
                                
                                if total_posts > 0:
                                    overall_progress = (total_completed / total_posts) * 100
                                    update_queue.put(("progress", min(100, overall_progress), 
                                                    f"#{hashtag}: {completed}/{total} posts"))
                            
                            # Call the scraper function
                            return get_scraper_data(hashtag, progress_callback)
                        except Exception as e:
                            error_holder["errors"].append(f"Error scraping #{hashtag}: {str(e)}")
                            return None
                    
                    # Start threads
                    threads = []
                    for hashtag in hashtags_to_process:
                        thread = Thread(
                            target=lambda h=hashtag: result_holder["data"].update({h: run_scraper(h)})
                        )
                        threads.append(thread)
                        thread.start()
                    
                    # Process updates
                    while any(t.is_alive() for t in threads):
                        try:
                            while True:
                                try:
                                    msg_type, value, message = update_queue.get_nowait()
                                    if msg_type == "progress":
                                        # Convert percentage to 0-1 range for Streamlit
                                        progress_bar.progress(value / 100)
                                        status_text.info(message)
                                except queue.Empty:
                                    break
                            time.sleep(0.1)
                        except Exception as e:
                            st.error(f"Error processing queue: {str(e)}")
                            break
                    
                    # Clean up threads
                    for thread in threads:
                        thread.join()
                    
                    # Handle results
                    if error_holder["errors"]:
                        for error in error_holder["errors"]:
                            st.error(error)
                    
                    # Store successful results
                    scraped_data = {
                        hashtag: result_holder["data"][hashtag] 
                        for hashtag in hashtags_to_process 
                        if result_holder["data"].get(hashtag)
                    }
                    
                    if scraped_data:
                        st.session_state.social_media_data = scraped_data
                        self.auth.update_project(
                                st.session_state.current_project["_id"],
                                {
                                    "social_media_data": scraped_data,
                                }
                            )
                        progress_bar.progress(1.0)
                        status_text.success(f"✅ Scraping completed for {len(scraped_data)} hashtags!")

        try:
            # Display results
            if 'social_media_data' in st.session_state and st.session_state.social_media_data or st.session_state.current_project.get('social_media_data'):
                social_data = st.session_state.get('social_media_data') or \
                (st.session_state.get('current_project') or {}).get('social_media_data')

                st.subheader("Scraping Results")
                
                # Prepare combined DataFrame for all hashtags
                all_data = []
                
                for hashtag, data in social_data.items():
                    with st.expander(f"Results for #{hashtag} ({len(data)} posts)", expanded=False):
                        if data:
                            # Process data for display (same as your original structure)
                            processed_data = []
                            for item in data:
                                processed_item = item.copy()
                                if isinstance(item.get("comments"), list):
                                    processed_item["comments"] = " | ".join(
                                        c.get("text", "") for c in item["comments"] if isinstance(c, dict))
                                processed_item["hashtag"] = hashtag  # Add hashtag column
                                processed_data.append(processed_item)
                            
                            # Create DataFrame for this hashtag
                            df = pd.json_normalize(processed_data)
                            st.dataframe(df, use_container_width=True)
                            
                            # Add to combined data
                            all_data.extend(processed_data)
                        else:
                            st.warning(f"No data found for #{hashtag}")
                
                # Create combined CSV for download
                if all_data:
                    combined_df = pd.json_normalize(all_data)
                    
                    # CSV download button (same as your original structure)
                    csv_buffer = io.StringIO()
                    combined_df.to_csv(csv_buffer, index=False)
                    st.download_button(
                        label="📥 Download All Data as CSV",
                        data=csv_buffer.getvalue(),
                        file_name="scraped_social_data.csv",
                        mime="text/csv"
                    )
        except:

            if st.session_state.get('social_media_data') or st.session_state.current_project.get('social_media_data'):
                st.markdown("### Scraped Data Preview")
                st.markdown(f"#### Last Hashtag {st.session_state.current_project.get('last_hashtags') or st.session_state.last_hashtags}")

                try:
                    data = st.session_state.current_project.get('social_media_data').copy() 
                except:
                    data = st.session_state.get('social_media_data').copy()

                # Process data for display
                for item in data:
                    if isinstance(item.get("comments"), list):
                        item["comments"] = " | ".join(
                            c.get("text", "") for c in item["comments"] if isinstance(c, dict))
                
                df = pd.json_normalize(data)
                st.dataframe(df, use_container_width=True)

                # CSV download
                csv_buffer = io.StringIO()
                df.to_csv(csv_buffer, index=False)
                st.download_button(
                    label="📥 Download as CSV",
                    data=csv_buffer.getvalue(),
                    file_name="scraped_social_data.csv",
                    mime="text/csv"
                )




        # Navigation buttons
        col1, col2 = st.columns([1, 1])
        with col1:
            if st.button("← Back to Projects", key="back_to_projects"):
                st.session_state.page = "projects"
                st.rerun()
        
        if 1 not in st.session_state.completed_steps:
            with col2:
                try:
                    st.session_state.social_media_data = st.session_state.current_project["social_media_data"]
                except:
                    pass
                
                if st.button("Start Digester →", key="start_Digester", 
                            disabled=not (st.session_state.rnd_files or st.session_state.mkt_files or 
                                        st.session_state.research_result or st.session_state.social_media_data)):
                    # Save files to Azure and update metadata
                    with st.spinner("Saving files to system..."):
                        tech_metadata = self.save_files_to_azure(st.session_state.rnd_files, "technical")
                        mkt_metadata = self.save_files_to_azure(st.session_state.mkt_files, "marketing")
                        
                        # Update file metadata in session and database
                        st.session_state.file_metadata = {
                            "technical": tech_metadata,
                            "marketing": mkt_metadata
                        }
                        self.auth.save_file_metadata(st.session_state.current_project["_id"], st.session_state.file_metadata)
                    
                    # Move to next step
                    st.session_state.completed_steps.append(1)
                    st.session_state.wizard_step = 2
                    st.rerun()



    # def run_agents(self):
    #     st.subheader("⚙️ Running Digester")
    #     st.info("This may take a few minutes. Please don't close your browser.")
        
    #     rnd_text = self.process_files(st.session_state.rnd_files)
    #     mkt_text = self.process_files(st.session_state.mkt_files)
        
    #     # Add research content to Digester
    #     if st.session_state.get('research_result'):
    #         rnd_text += "\n\nRESEARCH FINDINGS:\n" + str(st.session_state.research_result)
    #     if st.session_state.get("social_media_data"):
    #         mkt_text += "\n\nSOCIAL MEDIA CONTENT:\n" + str(st.session_state.social_media_data)
    #         rnd_text += "\n\nSOCIAL MEDIA CONTENT:\n" + str(st.session_state.social_media_data)
    #     if not rnd_text and not mkt_text:
    #         st.error("No valid text extracted from files")
    #         return
    #     # # # text combining 
    #     # text="Rnd text : "+rnd_text+" marketing text : "+mkt_text
    #     agents = {
    #         "IngredientsAgent": (run_ingredients, rnd_text),
    #         "TechnologyAgent": (run_technology, rnd_text),
    #         "BenefitsAgent": (run_benefits, rnd_text),
    #         "SituationsAgent": (run_situations, mkt_text),
    #         "MotivationsAgent": (run_motivations, mkt_text),
    #         "OutcomesAgent": (run_outcomes, mkt_text)
    #     }


    #     progress_bar = st.progress(0)
    #     results = {}
        
    #     with ThreadPoolExecutor() as executor:
    #         futures = {executor.submit(fn, text): name for name, (fn, text) in agents.items()}
    #         for i, future in enumerate(as_completed(futures)):
    #             name = futures[future]
    #             try:
    #                 result = future.result()
    #                 results[name] = result
    #                 st.success(f"✅ {name} completed")
    #             except Exception as e:
    #                 st.error(f"❌ {name} failed: {str(e)}")
    #             progress_bar.progress((i + 1) / len(agents))
        
    #     if results:
    #         self.auth.save_agent_results(st.session_state.current_project["_id"], results)
    #         st.session_state.agent_outputs = results
    #         st.session_state.completed_steps = [1, 2, 3]
    #         st.session_state.wizard_step = 3
    #         st.rerun()


    def load_lottie_url(self,url: str):
        r = requests.get(url)
        if r.status_code != 200:
            return None
        return r.json()

    def run_agents(self):


        # print("in run agents runner")
        # print(st.session_state.current_project["description"])
        # print(st.session_state.current_project["description"])
        # print(st.session_state.current_project["description"])
        # print(st.session_state.current_project["description"])
        # print(st.session_state.current_project["description"])
        # print(st.session_state.current_project["description"])
        # print(st.session_state.current_project["description"])
        # print(st.session_state.current_project["description"])
        # print(st.session_state.current_project["description"])
        st.subheader("⚙️ Running Digester")
        # time.sleep(200)
        # Create a visually appealing header
        st.markdown("""
        <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 20px;">
            <div style="font-size: 24px;">🔬</div>
            <div>
                <h3 style="margin: 0;">Research Digest in Progress</h3>
                <p style="margin: 0; color: #666;">Analyzing your documents with specialized agents</p>
            </div>
        </div>
        """, unsafe_allow_html=True)
     
        st.info("⏳ This may take a few minutes. Please don't close your browser.")
        status_text = st.empty()
        
        # Process files
        rnd_text = self.process_files(st.session_state.rnd_files)
        mkt_text = self.process_files(st.session_state.mkt_files)
        
        # Add research content to Digester
        if st.session_state.get('research_result'):
            rnd_text += "\n\nRESEARCH FINDINGS:\n" + str(st.session_state.research_result)
        if st.session_state.get("social_media_data"):
            mkt_text += "\n\nSOCIAL MEDIA CONTENT:\n" + str(st.session_state.social_media_data)
            rnd_text += "\n\nSOCIAL MEDIA CONTENT:\n" + str(st.session_state.social_media_data)
        if not rnd_text and not mkt_text:
            st.error("No valid text extracted from files")
            return
        project_description=st.session_state.current_project["description"]
        # Define agents
        agents = {
            "IngredientsAgent": (run_ingredients, rnd_text, project_description),
            "TechnologyAgent": (run_technology, rnd_text, project_description),
            "BenefitsAgent": (run_benefits, rnd_text, project_description),
            "SituationsAgent": (run_situations, mkt_text, project_description),
            "MotivationsAgent": (run_motivations, mkt_text, project_description),
            "OutcomesAgent": (run_outcomes, mkt_text, project_description),
        }
        # Create progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        results = {}
        agent_count = len(agents)
        
        # Create a container for agent status messages
        status_container = st.container()
        
        # Initialize status messages
        status_messages = {
            agent: status_container.empty() for agent in agents
        }
        
        # Set initial status
        for agent in agents:
            status_messages[agent].info(f"🟡 Starting...  {agent.replace('Agent', '')} ")
        
        # Run agents with ThreadPoolExecutor
        with ThreadPoolExecutor() as executor:
            futures = {executor.submit(fn, text,description): agent for agent, (fn, text,description) in agents.items()}
            
            for i, future in enumerate(as_completed(futures)):
                agent = futures[future]
                
                try:
                    result = future.result()
                    results[agent] = result
                    # Update status to success
                    status_messages[agent].success(f"✅ {agent.replace('Agent', '')} completed successfully")
                    
                    # Show quick notification
                    st.toast(f"✨ {agent.replace('Agent', '')} finished!", icon="✅")
                except Exception as e:
                    status_messages[agent].error(f"❌ {agent.replace('Agent', '')} failed: {str(e)}")
                    st.toast(f"⚠️ {agent.replace('Agent', '')} failed", icon="❌")
                
                # Update progress bar
                progress = int((i + 1) / agent_count * 100)
                progress_bar.progress(progress)
                status_text.info(f"⏳ Progress: {progress}% complete | {i+1}/{agent_count} agents finished")
        
        # Final completion
        progress_bar.progress(100)
        status_text.success("🎉 All agents completed!")
        
        # Celebration effect
        # st.balloons()
        st.success("Digester analysis complete! Ready to review results.")
        
        # Save results and move to next step
        if results:
            self.auth.save_agent_results(st.session_state.current_project["_id"], results)
            st.session_state.agent_outputs = results
            st.session_state.completed_steps = [1, 2, 3]
            st.session_state.wizard_step = 3
            
            # Auto-rerun after short delay to show results
            time.sleep(2)
            st.rerun()

    def normalize_agent_data(self, data):
        if isinstance(data, dict):
            for key in list(data.keys()):
                if key in ['Evidence_Snippets', 'Keywords', 'BenefitsAgent', 'TechnologyAgent']:
                    if not isinstance(data[key], list):
                        data[key] = [data[key]] if data[key] else []
                elif isinstance(data[key], dict):
                    data[key] = self.normalize_agent_data(data[key])
                elif isinstance(data[key], list):
                    data[key] = [self.normalize_agent_data(item) if isinstance(item, dict) else item 
                               for item in data[key]]
        return data

    def format_dataframe(self, data):
        try:
            clean_data = self.normalize_agent_data(data)
            if isinstance(clean_data, dict):
                df = json_normalize(clean_data, sep='_', errors='ignore')
            elif isinstance(clean_data, list):
                df = json_normalize(clean_data, sep='_', errors='ignore')
            else:
                df = pd.DataFrame([clean_data])
            
            return df.astype(str).replace({
                'nan': '', 'None': '', 'NaT': '', '[]': '[]', '{}': '{}'
            })
        except Exception as e:
            st.error(f"📊 Formatting error: {str(e)}")
            return pd.DataFrame({"Raw Data": [str(clean_data)]})



  

    def show_results(self):
        st.subheader("📊 Digester Results")
        
        if not st.session_state.agent_outputs:
            st.warning("No Digester results found")
            return

        # Initialize session state
        if "selected_rows" not in st.session_state:
            st.session_state.selected_rows = {}
        if "select_all_states" not in st.session_state:
            st.session_state.select_all_states = {}

        # Track current selections
        current_selections_exist = False

        for agent_name, agent_data in st.session_state.agent_outputs.items():
            if agent_name == "ProductGenerationAgent":
                continue

            df = self.format_dataframe(agent_data)
            
            if not df.empty:
                # Initialize agent's select all state
                if agent_name not in st.session_state.select_all_states:
                    st.session_state.select_all_states[agent_name] = True
                
                # Add Select column initialized with select all state
                df.insert(0, "Select", st.session_state.select_all_states[agent_name])
                
                with st.expander(f"🔎 {agent_name.replace('Agent', ' Digester')}",expanded=False):
                    # Select all checkbox with synchronization
                    select_all = st.checkbox(
                        "Select all",
                        value=st.session_state.select_all_states[agent_name],
                        key=f"select_all_{agent_name}",
                        on_change=lambda a=agent_name: self._handle_select_all(a)
                    )
                    
                    # Display the dataframe
                    edited_df = st.data_editor(
                        df,
                        key=f"editor_{agent_name}",
                        use_container_width=True,
                        disabled=df.columns[1:],  # Only allow editing Select column
                        hide_index=True
                    )
                    
                    # Update selection states
                    selected_rows = edited_df[edited_df["Select"]]
                    if not selected_rows.empty:
                        st.session_state.selected_rows[agent_name] = selected_rows.to_dict('records')
                        current_selections_exist = True
                        # Update select all state if all rows are selected
                        if len(selected_rows) == len(df):
                            st.session_state.select_all_states[agent_name] = True
                        else:
                            st.session_state.select_all_states[agent_name] = False
                    elif agent_name in st.session_state.selected_rows:
                        del st.session_state.selected_rows[agent_name]
                        st.session_state.select_all_states[agent_name] = False

                    # Show raw JSON option
                    if st.checkbox(f"Show raw JSON for {agent_name}", key=f"raw_{agent_name}"):
                        st.json(agent_data)

        
        st.info("You can select the specific raws from the above tables to proceed with the study from results directly")
        # Navigation buttons
        st.markdown("---")
         # Navigation buttons
        col1, col2 = st.columns([1,1])
        with col1:
            if st.button("← Back to File Upload", key="back_to_step1"):
                st.session_state.wizard_step = 1
                st.rerun()

        with col2:
            # Check if product ideas exist
            has_products = "ProductGenerationAgent" in st.session_state.agent_outputs
            has_generations = has_products and st.session_state.agent_outputs["ProductGenerationAgent"].get("generations")
            
            if has_generations:
                if st.button("View Product Ideas →", key="view_products"):
                    st.session_state.wizard_step = 4
                    st.rerun()
            else:
                if st.button("Generate Product Ideas →", key="generate_products"):
                    with st.spinner("Generating product ideas..."):
                        if self.generate_products():
                            st.rerun()
    
        # Clear selections handler
        if "clear_clicked" not in st.session_state:
            st.session_state.clear_clicked = False

        # Show action buttons only if we have current selections
        if current_selections_exist or (st.session_state.selected_rows and not st.session_state.clear_clicked):
            st.markdown("---")
            st.info(f"You have selected rows from the following digesters :- {', '.join([k.replace('Agent', ' Digester') for k in list(st.session_state.selected_rows.keys())])}")
            col1, col2 = st.columns(2)
            with col1:
                
                if st.button("📚 Generate Study", type="primary", key=f"gen-study"):
                    st.info("We are working of this functionality...")
                            # st.session_state["study_step"] = 0
                            # st.session_state["selected_idea_idx"] =  1  # store 0-based index

            with col2:
                if st.button("❌ Clear Selections"):
                    st.session_state.selected_rows = {}
                    st.session_state.select_all_states = {k: False for k in st.session_state.select_all_states}
                    st.session_state.clear_clicked = True
                    st.rerun()
            # # Now handle study generation AFTER loop (full-width)
            # if st.session_state.get("study_step", 0) >= 0:
            #     # selected_idx = st.session_state.get("selected_idea_idx", 0)
            #     selected_idea = st.session_state.selected_rows
            #     from .study import StudyGenerationProcess
            #     study_gen = StudyGenerationProcess(self.auth, selected_idea,type="results")
            #     study_gen.run()
        # # Reset clear flag after render
        # if st.session_state.clear_clicked:
        #     st.session_state.clear_clicked = False
    def _handle_select_all(self, agent_name):
        """Handle select all checkbox changes"""
        # Toggle the select all state
        st.session_state.select_all_states[agent_name] = not st.session_state.select_all_states.get(agent_name, False)
        
        # Update the dataframe editor state
        if f"editor_{agent_name}" in st.session_state:
            edited_data = st.session_state[f"editor_{agent_name}"]["edited_rows"]
            for row_idx in edited_data:
                edited_data[row_idx]["Select"] = st.session_state.select_all_states[agent_name]




    def generate_product_docx(self, product_data):
        """Generate a Word document for a product idea"""
        doc = Document()
        
        # Add title
        doc.add_heading(product_data.get('product_name', 'Product Idea'), level=1)
        
        # Add sections
        sections = [
            ('Technical Explanation', product_data.get('technical_explanation')),
            ('Consumer Pitch', product_data.get('consumer_pitch')),
            ('Competitor Reaction', product_data.get('competitor_reaction')),
            ('5-Year Projection', product_data.get('five_year_projection')),
            ('Consumer Discussion', product_data.get('consumer_discussion')),
            ('Professional Presentation', '\n'.join([
                f"{i+1}. {item}" for i, item in enumerate(product_data.get('presentation', []))
            ])),
            ('Investor Evaluation', product_data.get('investor_evaluation'))
        ]
        
        for title, content in sections:
            if content:
                doc.add_heading(title, level=2)
                doc.add_paragraph(content)
        
        # Add Q&A
        qa = product_data.get('consumer_qa', [])
        if qa:
            doc.add_heading('Consumer Q&A', level=2)
            for i, item in enumerate(qa, 1):
                doc.add_paragraph(f"Q{i}: {item.get('question', '')}", style='Heading3')
                for j, answer in enumerate(item.get('answers', []), 1):
                    doc.add_paragraph(f"{j}. {answer}")
        
        # Add slogans
        slogans = product_data.get('advertisor_slogans', [])
        if slogans:
            doc.add_heading('Advertiser Slogans', level=2)
            for slogan in slogans:
                doc.add_paragraph(f"Slogan: {slogan.get('slogan', '')}")
                doc.add_paragraph(f"Mindset: {slogan.get('mindset_description', '')}")
        
        # Create in-memory file
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream


    def generate_products(self):
        """Run product generation with version history"""
        st.subheader("🚀 Generating Product Ideas")
        
        # Get all agent outputs
        all_agent_outputs = st.session_state.agent_outputs
        # if st.session_state.selected_rows=={}:
        #     all_agent_outputs = st.session_state.agent_outputs
        # else:
        #     # selected_raws=list(st.session_state.selected_rows.keys())
        #     # for i in selected_raws:
        #     all_agent_outputs = st.session_state.selected_rows
        # print(all_agent_outputs)
        if not all_agent_outputs:
            st.error("No Digester results found. Please run Digester first.")
            return False

        # Create progress container
        progress_container = st.container()
        progress_bar = progress_container.progress(0)
        status_text = progress_container.empty()
        
        def progress_callback(percent, message):
            progress_bar.progress(percent)
            status_text.text(message)
        project_description=st.session_state.current_project["description"]
        try:
            # Run product generation agent
            product_ideas = run_product_generation(all_agent_outputs,project_description, progress_callback=progress_callback)
            
            # Store the results in session
            if "ProductGenerationAgent" not in st.session_state.agent_outputs:
                st.session_state.agent_outputs["ProductGenerationAgent"] = {
                    "generations": [],
                    "current_generation": None
                }
            
            # Create new generation
            new_gen = {
                "_id":  str(uuid.uuid4()),
                "timestamp": datetime.now().isoformat(),
                "ideas": product_ideas
            }
            
            # Add to history
            st.session_state.agent_outputs["ProductGenerationAgent"]["generations"].append(new_gen)
            st.session_state.agent_outputs["ProductGenerationAgent"]["current_generation"] = new_gen
            
            # Save to database
            self.auth.save_product_results(
                st.session_state.current_project["_id"],
                product_ideas
            )
            
            # Update UI state
            st.session_state.completed_steps = [1, 2, 3, 4]
            st.session_state.wizard_step = 4
            
            # Refresh the UI
            st.rerun()
            return True
            
        except Exception as e:
            progress_callback(100, f"Error: {str(e)}")
            time.sleep(2)
            return False
        finally:
            # Clear progress after delay
            time.sleep(1)
            progress_container.empty()



    def show_product_ideas(self):
        st.subheader("💡 Generated Product Ideas")
        
        # Check if we have product ideas
        if "ProductGenerationAgent" not in st.session_state.agent_outputs:
            st.warning("No product ideas generated yet")
            return

        product_data = st.session_state.agent_outputs["ProductGenerationAgent"]
        
        # Handle old format
        if isinstance(product_data, list):
            generations = [{
                "_id":  str(uuid.uuid4()),
                "created_at": datetime.now().isoformat(),
                "ideas": product_data
            }]
            current_generation = 0
        else:
            generations = product_data.get("generations", [])
            current_generation = product_data.get("current_generation")
        
        if not generations:
            st.warning("No product ideas generated yet")
            return
        
        # Generation selector
        st.markdown("### Product Generation Versions")
        gen_options = []
        for i, gen in enumerate(generations):
            # Handle different timestamp formats
            created_at = gen.get("created_at")
            if isinstance(created_at, str):
                try:
                    dt = datetime.fromisoformat(created_at)
                    time_str = dt.strftime("%Y-%m-%d %H:%M")
                except:
                    time_str = created_at
            elif isinstance(created_at, datetime):
                time_str = created_at.strftime("%Y-%m-%d %H:%M")
            else:
                time_str = "Unknown time"
                
            gen_options.append(f"Generation #{i+1} ")
        
        # Add "Create New" option
        gen_options.append("➕ Create New Generation")
        
        # Default to latest generation
        default_idx = len(generations) - 1
        
        # Create selector
        selected_idx = st.selectbox(
            "Select product generation:",
            options=range(len(gen_options)),
            index=default_idx,
            format_func=lambda i: gen_options[i]
        )
        
        # Handle "Create New" selection
        if selected_idx == len(generations):
            if st.button("🔄 Generate New Product Set", use_container_width=True):
                with st.spinner("Generating new product ideas..."):
                    if self.generate_products():
                        st.rerun()
            return
        
        # Get selected generation data
        selected_gen = generations[selected_idx]
        ideas = selected_gen["ideas"]
        # If we have a string, try to parse it as JSON
        if isinstance(ideas, str):
            try:
                ideas = json.loads(ideas)
            except json.JSONDecodeError:
                st.error("Could not parse product ideas as JSON")
                st.text_area("Raw Output", ideas, height=400)
                return
        
        # Display each product idea in a structured format
        if isinstance(ideas, list) and ideas:
            for idx, idea in enumerate(ideas, 1):
                cols = st.columns([0.85, 0.15]) 
                with cols[0]:
                    with st.expander(f"Idea #{idx}: {idea.get('product_name', 'Unnamed Product')}", expanded=False):
                        # Product header
                        st.markdown(f"### {idea.get('product_name', 'Unnamed Product')}")
                        
                        # Technical Explanation
                        st.markdown("#### Technical Explanation")
                        st.write(idea.get("technical_explanation", "No technical explanation available"))
                        
                        # Consumer Pitch
                        st.markdown("#### Consumer Pitch")
                        st.write(idea.get("consumer_pitch", "No consumer pitch available"))
                        
                        # Competitor Reaction
                        st.markdown("#### Competitor Reaction")
                        st.write(idea.get("competitor_reaction", "No competitor reaction available"))
                        
                        # 5-Year Projection
                        st.markdown("#### 5-Year Projection (2030)")
                        st.write(idea.get("five_year_projection", "No projection available"))
                        
                        # Consumer Discussion
                        st.markdown("#### Consumer Discussion (Town Hall Meeting)")
                        st.write(idea.get("consumer_discussion", "No consumer discussion available"))
                        
                        # Presentation
                        st.markdown("#### Professional Presentation")
                        presentation = idea.get("presentation", [])
                        if isinstance(presentation, list) and len(presentation) > 0:
                            for i, sentence in enumerate(presentation, 1):
                                st.markdown(f"{i}. {sentence}")
                        else:
                            st.write("No presentation available")
                        
                        # Consumer Q&A
                        st.markdown("#### Consumer Q&A")
                        qa = idea.get("consumer_qa", [])
                        if isinstance(qa, list) and len(qa) > 0:
                            for i, qa_item in enumerate(qa, 1):
                                st.markdown(f"**Q{i}:** {qa_item.get('question', '')}")
                                answers = qa_item.get('answers', [])
                                if answers:
                                    for j, ans in enumerate(answers, 1):
                                        st.markdown(f"{j}. {ans}")
                                else:
                                    st.markdown("No answers available")
                        else:
                            st.write("No Q&A available")
                        
                        # Investor Evaluation
                        st.markdown("#### Investor Evaluation")
                        st.write(idea.get("investor_evaluation", "No investor evaluation available"))
                        
                        # Advertiser Slogans
                        st.markdown("#### Advertiser Slogans")
                        slogans = idea.get("advertisor_slogans", [])
                        if isinstance(slogans, list) and len(slogans) > 0:
                            for slogan in slogans:
                                st.markdown(f"**Slogan:** {slogan.get('slogan', '')}")
                                st.markdown(f"**Mindset Description:** {slogan.get('mindset_description', '')}")
                                st.markdown("---")
                        else:
                            st.write("No slogans available")
                        
                        # AI Report Card
                        st.markdown("#### AI Report Card")
                        report_card = idea.get("ai_report_card", {})
                        if report_card:
                            report_data = {
                                "Metric": ["Originality", "Usefulness", "Social Media Talk", 
                                        "Memorability", "Friend Talk", "Purchase Ease", 
                                        "Excitement", "Boredom Likelihood"],
                                "Score": [
                                    report_card.get("originality", 0),
                                    report_card.get("usefulness", 0),
                                    report_card.get("social_media_talk", 0),
                                    report_card.get("memorability", 0),
                                    report_card.get("friend_talk", 0),
                                    report_card.get("purchase_ease", 0),
                                    report_card.get("excitement", 0),
                                    report_card.get("boredom_likelihood", 0)
                                ]
                            }
                            st.dataframe(report_data)
                        else:
                            st.write("No AI report card available")
                        
                        st.markdown("---")
                    with cols[1]:
                        if st.button("📚 Prepare Study", type="primary", key=f"gen-study-{idx}"):
                            st.session_state["study_step"] = 0
                            st.session_state["selected_idea_idx"] = idx - 1  # store 0-based index

        # Now handle study generation AFTER loop (full-width)
        if st.session_state.get("study_step", 0) >= 0:
            selected_idx = st.session_state.get("selected_idea_idx", 0)
            selected_idea = ideas[selected_idx]
            from .study import StudyGenerationProcess
            study_gen = StudyGenerationProcess(self.auth, selected_idea)
            study_gen.run()
                

            # Add JSON viewer at the bottom
            # with st.expander("View Raw JSON Output", expanded=False):
            #     st.json(ideas)

            
        # Navigation buttons
        col1, col2 = st.columns([1,1])
        with col1:
            if st.button("← Back to Results", key="back_to_results"):
                st.session_state.wizard_step = 3
                st.rerun()
        with col2:
            if st.button("🔄 Regenerate Products", key="regenerate_products"):
                with st.spinner("Regenerating product ideas..."):
                    if self.generate_products():
                        st.rerun()